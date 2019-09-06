#!/bin/python
from docx.enum.style import WD_STYLE_TYPE
from docx import Document

doc = Document()

# doc.add_heading('测试语句',level = 0)

# doc.add_paragraph('第二行','Subtitle')
# doc.add_paragraph('zhengwen zhengwen')
# doc.save('testdoc.docx')

# for s in doc.styles:
#     if s.type == WD_STYLE_TYPE.PARAGRAPH:
#         print(s)
#         print(s.font)
#         print(s.name)

p = doc.add_paragraph('222222',style = 'Title')
p.add_run('333runrunrun')

para = doc.add_paragraph('putong duanluo')
para.add_run('add txt')
para.add_run('add doc')

from docx.enum.text import WD_ALIGN_PARAGRAPH
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para_left = doc.add_paragraph('左对齐')

para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
left_run = para_left.add_run('单独调整样式')


from docx.shared import Pt
left_run.font.bold = True
left_run.font.size = Pt(20)

doc.add_picture('印章.png')




table = doc.add_table(rows=5,cols=8)

cell = table.cell(1,1)
cell.text = 'ufo'
table.style = 'Light Grid Accent 5'
table.cell(3,0).text = "追加"
doc.save('testdoc.docx')
